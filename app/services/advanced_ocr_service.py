from __future__ import annotations

import csv
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from flask import current_app

from app.models import ManagedFile
from app.services.ocr_service import OCRService
from app.services.pdf_service import PDFService
from app.services.storage_service import StorageService


class AdvancedOCRService:
    @staticmethod
    def _preprocess_image(source_path: Path, output_path: Path, options: dict) -> Path:
        try:
            import cv2  # type: ignore
            import numpy as np  # type: ignore
        except Exception as exc:
            raise ValueError("OpenCV and numpy are required for OCR preprocessing.") from exc

        frame = cv2.imread(str(source_path))
        if frame is None:
            raise ValueError("Unable to read image for OCR.")

        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)

        if options.get("denoise"):
            gray = cv2.fastNlMeansDenoising(gray, None, 15, 7, 21)

        if options.get("contrast_boost"):
            gray = cv2.convertScaleAbs(gray, alpha=1.35, beta=8)

        if options.get("black_white"):
            _, gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)

        if options.get("deskew"):
            inverted = cv2.bitwise_not(gray)
            _, threshold = cv2.threshold(inverted, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
            coordinates = np.column_stack(np.where(threshold > 0))
            if coordinates.size:
                angle = cv2.minAreaRect(coordinates)[-1]
                if angle < -45:
                    angle = -(90 + angle)
                else:
                    angle = -angle
                height, width = gray.shape[:2]
                matrix = cv2.getRotationMatrix2D((width / 2, height / 2), angle, 1.0)
                gray = cv2.warpAffine(
                    gray,
                    matrix,
                    (width, height),
                    flags=cv2.INTER_CUBIC,
                    borderMode=cv2.BORDER_REPLICATE,
                )

        output_path.parent.mkdir(parents=True, exist_ok=True)
        cv2.imwrite(str(output_path), gray)
        return output_path

    @staticmethod
    def _write_docx(text_path: Path, output_path: Path) -> Path:
        from docx import Document

        document = Document()
        document.add_heading("OCR Output", level=1)
        content = text_path.read_text(encoding="utf-8", errors="ignore")
        for paragraph in content.split("\n\n"):
            if paragraph.strip():
                document.add_paragraph(paragraph.strip())
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
        return output_path

    @staticmethod
    def _extract_tables(pdf_path: Path, csv_path: Path, xlsx_path: Path) -> tuple[Path, Path]:
        import pdfplumber
        from openpyxl import Workbook

        rows: list[list[str]] = []
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "OCR Tables"
        output_row = 1

        with pdfplumber.open(str(pdf_path)) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                tables = page.extract_tables() or []
                if not tables:
                    continue
                for table in tables:
                    for row in table:
                        values = [str(cell or "").strip() for cell in (row or [])]
                        if not any(values):
                            continue
                        rows.append(values)
                        for column_index, value in enumerate(values, start=1):
                            sheet.cell(row=output_row, column=column_index, value=value)
                        output_row += 1
                    rows.append([])
                    output_row += 1

        csv_path.parent.mkdir(parents=True, exist_ok=True)
        with csv_path.open("w", encoding="utf-8", newline="") as handle:
            writer = csv.writer(handle)
            for row in rows:
                writer.writerow(row)

        workbook.save(str(xlsx_path))
        return csv_path, xlsx_path

    @staticmethod
    def _estimate_confidence(image_path: Path, lang: str) -> float | None:
        try:
            pytesseract = OCRService._tesseract()  # noqa: SLF001
            from PIL import Image
        except Exception:
            return None

        data = pytesseract.image_to_data(Image.open(image_path), lang=lang, output_type=pytesseract.Output.DICT)
        values = []
        for conf in data.get("conf", []):
            try:
                value = float(conf)
            except Exception:
                continue
            if value >= 0:
                values.append(value)
        if not values:
            return None
        return round(sum(values) / len(values), 2)

    @staticmethod
    def process_upload(
        *,
        user_id: int,
        file_record: ManagedFile,
        options: dict,
    ) -> dict:
        source_path = StorageService.absolute_path(file_record)
        suffix = source_path.suffix.lower()
        if suffix not in {".pdf", ".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}:
            raise ValueError("OCR Suite supports PDF and image files only.")

        languages = (options.get("languages") or "eng").strip().replace(" ", "")
        if not languages:
            languages = "eng"

        requested_formats = options.get("output_formats") or ["txt"]
        if isinstance(requested_formats, str):
            requested_formats = [fmt.strip().lower() for fmt in requested_formats.split(",") if fmt.strip()]
        requested_formats = sorted(set(fmt.lower() for fmt in requested_formats if fmt))
        if not requested_formats:
            requested_formats = ["txt"]

        work_dir = Path(current_app.config["OUTPUT_ROOT"]) / str(user_id) / "advanced_ocr" / file_record.stored_name
        work_dir.mkdir(parents=True, exist_ok=True)

        produced_paths: list[Path] = []
        text_path = work_dir / f"{source_path.stem}_ocr.txt"
        searchable_pdf_path = work_dir / f"{source_path.stem}_searchable.pdf"
        confidence_value: float | None = None

        if suffix == ".pdf":
            OCRService.ocr_pdf_to_searchable(
                input_pdf_path=str(source_path),
                output_pdf_path=str(searchable_pdf_path),
                output_text_path=str(text_path),
                lang=languages,
            )
            if "searchable_pdf" in requested_formats:
                produced_paths.append(searchable_pdf_path)
            if "txt" in requested_formats:
                produced_paths.append(text_path)

            preview_image_dir = work_dir / "preview"
            preview_images = PDFService.pdf_to_images(str(source_path), str(preview_image_dir), image_format="png")
            if preview_images:
                confidence_value = AdvancedOCRService._estimate_confidence(Path(preview_images[0]), languages)

            if any(fmt in requested_formats for fmt in {"csv", "xlsx"}):
                csv_path, xlsx_path = AdvancedOCRService._extract_tables(
                    source_path,
                    work_dir / f"{source_path.stem}_tables.csv",
                    work_dir / f"{source_path.stem}_tables.xlsx",
                )
                if "csv" in requested_formats:
                    produced_paths.append(csv_path)
                if "xlsx" in requested_formats:
                    produced_paths.append(xlsx_path)

        else:
            preprocessed_path = work_dir / f"{source_path.stem}_preprocessed.png"
            if any(options.get(flag) for flag in ("deskew", "denoise", "contrast_boost", "black_white")):
                preprocessed_path = AdvancedOCRService._preprocess_image(source_path, preprocessed_path, options)
                image_for_ocr = preprocessed_path
            else:
                image_for_ocr = source_path

            OCRService.ocr_image_paths([str(image_for_ocr)], str(text_path), lang=languages)
            if "txt" in requested_formats:
                produced_paths.append(text_path)

            if "searchable_pdf" in requested_formats:
                temp_pdf = work_dir / f"{source_path.stem}_scan.pdf"
                PDFService.images_to_pdf([str(image_for_ocr)], str(temp_pdf))
                OCRService.ocr_pdf_to_searchable(
                    input_pdf_path=str(temp_pdf),
                    output_pdf_path=str(searchable_pdf_path),
                    output_text_path=str(text_path),
                    lang=languages,
                )
                produced_paths.append(searchable_pdf_path)

            confidence_value = AdvancedOCRService._estimate_confidence(Path(image_for_ocr), languages)

        if "docx" in requested_formats:
            docx_path = work_dir / f"{source_path.stem}_ocr.docx"
            AdvancedOCRService._write_docx(text_path, docx_path)
            produced_paths.append(docx_path)

        produced_paths = [path for path in produced_paths if path.exists()]
        if not produced_paths:
            raise ValueError("No OCR outputs were generated.")

        if len(produced_paths) == 1:
            output_file = StorageService.register_existing_file(
                absolute_path=produced_paths[0],
                user_id=user_id,
                kind="output",
                original_name=produced_paths[0].name,
                label="Advanced OCR output",
            )
        else:
            bundle_path = work_dir / f"{source_path.stem}_ocr_bundle.zip"
            with ZipFile(bundle_path, "w", compression=ZIP_DEFLATED) as archive:
                for output_path in produced_paths:
                    archive.write(output_path, arcname=output_path.name)
            output_file = StorageService.register_existing_file(
                absolute_path=bundle_path,
                user_id=user_id,
                kind="output",
                original_name=bundle_path.name,
                label="Advanced OCR bundle",
            )

        return {
            "output_file": output_file,
            "output_count": len(produced_paths),
            "confidence": confidence_value,
            "languages": languages,
            "formats": requested_formats,
        }
